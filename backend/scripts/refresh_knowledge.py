"""Knowledge base refresh script with APScheduler integration.

Usage:
    uv run python scripts/refresh_knowledge.py --run-now
"""

import argparse
import asyncio
import logging
import sys
from pathlib import Path

# Ensure the backend root is on sys.path when running as a script
sys.path.insert(0, str(Path(__file__).parent.parent))

import yfinance as yf
from langchain_chroma import Chroma
from langchain_community.document_loaders import WebBaseLoader, WikipediaLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.settings import settings
from services.embedding import get_embeddings

logger = logging.getLogger(__name__)

# ── Knowledge source configuration ────────────────────────────
SOURCES = {
    # 使用对程序化抓取更友好的金融教育网站（替代 Investopedia，后者有反爬限制）
    "static_pages": [
        # CFI (Corporate Finance Institute) — 服务端渲染，抓取友好
        "https://corporatefinanceinstitute.com/resources/valuation/price-earnings-ratio/",
        "https://corporatefinanceinstitute.com/resources/accounting/earnings-per-share-eps-formula/",
        "https://corporatefinanceinstitute.com/resources/valuation/fcf-formula-free-cash-flow/",
        "https://corporatefinanceinstitute.com/resources/valuation/ebitda-margin/",
        # SEC Investor.gov — 美国政府官方投资者教育网站
        "https://www.investor.gov/introduction-investing/investing-basics/investment-products",
        "https://www.investor.gov/introduction-investing/investing-basics/how-stock-markets-work",
    ],
    "wikipedia": [
        "市盈率",
        "净资产收益率",
        "自由现金流",
        "市净率",
        "Price–earnings ratio",
        "EBITDA",
    ],
    "earnings_tickers": ["BABA", "TSLA", "AAPL", "TCEHY"],
    "tavily_queries": [
        "阿里巴巴最新季度财报摘要",
        "特斯拉季报要点",
        "腾讯控股财务数据",
    ],
}

splitter = RecursiveCharacterTextSplitter(
    chunk_size=512,
    chunk_overlap=64,
    separators=["\n\n", "\n", "。", ".", " "],
)

def _get_vectordb():
    """Lazy-init vectordb so embedding config is read at call time, not import time."""
    return Chroma(
        collection_name="finance_knowledge",
        embedding_function=get_embeddings(),
        persist_directory=settings.CHROMA_DIR,
    )


def _load_local_files() -> list[Document]:
    """Recursively load knowledge files (txt, md, docx, pdf) from KNOWLEDGE_FILES_DIR."""
    if not settings.KNOWLEDGE_FILES_ENABLED or not settings.KNOWLEDGE_FILES_DIR:
        logger.info("Local file loading skipped (disabled or no directory configured)")
        return []

    knowledge_dir = Path(settings.KNOWLEDGE_FILES_DIR)
    if not knowledge_dir.exists() or not knowledge_dir.is_dir():
        logger.warning(f"Knowledge files directory does not exist: {knowledge_dir}")
        return []

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
    docs: list[Document] = []
    file_count = 0

    for file_path in sorted(knowledge_dir.rglob("*")):
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        file_count += 1
        try:
            ext = file_path.suffix.lower()

            if ext in (".txt", ".md"):
                text = file_path.read_text(encoding="utf-8", errors="replace")
                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": str(file_path),
                            "type": "local_file",
                            "format": ext.lstrip("."),
                            "filename": file_path.name,
                        },
                    )
                )

            elif ext == ".docx":
                try:
                    import docx

                    doc = docx.Document(str(file_path))
                    text = "\n".join(
                        para.text for para in doc.paragraphs if para.text.strip()
                    )
                    if text.strip():
                        docs.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": str(file_path),
                                    "type": "local_file",
                                    "format": "docx",
                                    "filename": file_path.name,
                                },
                            )
                        )
                except ImportError:
                    logger.error(
                        "python-docx not installed, skipping .docx files. "
                        "Run: uv add python-docx"
                    )

            elif ext == ".pdf":
                try:
                    import fitz  # pymupdf

                    pdf_doc = fitz.open(str(file_path))
                    pages_text = []
                    for page in pdf_doc:
                        page_text = page.get_text()
                        if page_text.strip():
                            pages_text.append(page_text)
                    pdf_doc.close()

                    if pages_text:
                        text = "\n\n".join(pages_text)
                        docs.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": str(file_path),
                                    "type": "local_file",
                                    "format": "pdf",
                                    "filename": file_path.name,
                                    "page_count": len(pages_text),
                                },
                            )
                        )
                except ImportError:
                    logger.error(
                        "pymupdf not installed, skipping .pdf files. "
                        "Run: uv add pymupdf"
                    )

            logger.info(f"Loaded local file: {file_path.name}")

        except Exception as e:
            logger.error(f"Failed to load local file {file_path}: {e}")

    logger.info(
        f"Local file scan complete: {file_count} files found, "
        f"{len(docs)} documents loaded from {knowledge_dir}"
    )
    return docs


def _load_static_pages() -> list[Document]:
    """Load financial concept pages from Investopedia."""
    docs: list[Document] = []
    for url in SOURCES["static_pages"]:
        try:
            docs.extend(WebBaseLoader(url).load())
            logger.info(f"Loaded page: {url}")
        except Exception as e:
            logger.error(f"Failed to load page {url}: {e}")
    return docs


def _load_wikipedia() -> list[Document]:
    """Load Wikipedia articles for financial terms."""
    docs: list[Document] = []
    for q in SOURCES["wikipedia"]:
        try:
            lang = "zh" if any("\u4e00" <= c <= "\u9fff" for c in q) else "en"
            docs.extend(WikipediaLoader(query=q, load_max_docs=1, lang=lang).load())
            logger.info(f"Loaded Wikipedia: {q}")
        except Exception as e:
            logger.error(f"Wikipedia failed for {q}: {e}")
    return docs


def _load_earnings() -> list[Document]:
    """Load financial summaries from Yahoo Finance."""
    docs: list[Document] = []
    for ticker in SOURCES["earnings_tickers"]:
        try:
            info = yf.Ticker(ticker).info
            text = (
                f"# {ticker} 财务摘要（来源：Yahoo Finance）\n"
                f"市盈率（P/E）: {info.get('trailingPE', 'N/A')}\n"
                f"市净率（P/B）: {info.get('priceToBook', 'N/A')}\n"
                f"EPS: {info.get('trailingEps', 'N/A')}\n"
                f"营收（TTM）: {info.get('totalRevenue', 'N/A')}\n"
                f"净利润率: {info.get('profitMargins', 'N/A')}\n"
                f"ROE: {info.get('returnOnEquity', 'N/A')}\n"
                f"行业: {info.get('industry', 'N/A')}\n"
            )
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": f"yfinance:{ticker}", "type": "earnings"},
                )
            )
            logger.info(f"Loaded earnings: {ticker}")
        except Exception as e:
            logger.error(f"yfinance failed for {ticker}: {e}")
    return docs


def _load_tavily() -> list[Document]:
    """Load web search results via Tavily."""
    if not settings.TAVILY_API_KEY:
        logger.warning("Tavily API key not configured, skipping web sources")
        return []

    docs: list[Document] = []
    try:
        from tavily import TavilyClient

        tavily = TavilyClient(api_key=settings.TAVILY_API_KEY)

        for q in SOURCES["tavily_queries"]:
            try:
                for r in tavily.search(q, max_results=3).get("results", []):
                    docs.append(
                        Document(
                            page_content=r["content"],
                            metadata={"source": r["url"], "type": "web", "query": q},
                        )
                    )
                logger.info(f"Loaded Tavily results for: {q}")
            except Exception as e:
                logger.error(f"Tavily failed for {q}: {e}")
    except Exception as e:
        logger.error(f"Tavily client init failed: {e}")
    return docs


def refresh_knowledge_base():
    """Full refresh of the knowledge base. Called by APScheduler or manually."""
    logger.info("=== Knowledge base refresh started ===")

    all_docs = (
        _load_local_files()
        + _load_static_pages()
        + _load_wikipedia()
        + _load_earnings()
        + _load_tavily()
    )

    if not all_docs:
        logger.error("No documents retrieved, skipping this refresh cycle")
        return

    chunks = splitter.split_documents(all_docs)
    logger.info(f"Documents: {len(all_docs)}, Chunks: {len(chunks)}")

    vectordb = _get_vectordb()

    # Full rebuild (simple and reliable, avoids dedup complexity)
    try:
        vectordb.delete_collection()
    except Exception:
        pass  # Collection may not exist yet

    # Re-create vectordb instance after deletion, because delete_collection()
    # invalidates the internal collection reference in langchain-chroma
    vectordb = _get_vectordb()
    vectordb.add_documents(chunks)
    logger.info(f"=== Knowledge base refresh complete, wrote {len(chunks)} chunks ===")

    _log_refresh(len(all_docs), len(chunks))


def _log_refresh(doc_count: int, chunk_count: int):
    """Log refresh metadata to SQLite."""
    import aiosqlite

    async def _write():
        async with aiosqlite.connect(settings.SQLITE_PATH) as db:
            await db.execute(
                "CREATE TABLE IF NOT EXISTS kb_refresh_log "
                "(id INTEGER PRIMARY KEY, doc_count INT, chunk_count INT, refreshed_at TEXT)"
            )
            await db.execute(
                "INSERT INTO kb_refresh_log (doc_count, chunk_count, refreshed_at) "
                "VALUES (?, ?, datetime('now'))",
                (doc_count, chunk_count),
            )
            await db.commit()

    asyncio.run(_write())


def start_scheduler():
    """Register scheduled jobs. Called from FastAPI lifespan."""
    from apscheduler.schedulers.background import BackgroundScheduler

    scheduler = BackgroundScheduler(timezone="Asia/Shanghai")
    scheduler.add_job(
        refresh_knowledge_base,
        "cron",
        hour=settings.KB_REFRESH_CRON_HOUR,
        minute=0,
        id="kb_daily_refresh",
        replace_existing=True,
    )
    # Extra Monday refresh to cover weekend earnings releases
    scheduler.add_job(
        refresh_knowledge_base,
        "cron",
        day_of_week="mon",
        hour=3,
        minute=0,
        id="kb_weekly_refresh",
        replace_existing=True,
    )
    scheduler.start()
    logger.info("APScheduler started — knowledge base refresh jobs registered")
    return scheduler


# ── CLI manual trigger ────────────────────────────────────────
if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
    )
    parser = argparse.ArgumentParser(description="Knowledge base refresh utility")
    parser.add_argument(
        "--run-now", action="store_true", help="Immediately run a full refresh"
    )
    args = parser.parse_args()
    if args.run_now:
        refresh_knowledge_base()
    else:
        print("Use --run-now to trigger an immediate refresh")
