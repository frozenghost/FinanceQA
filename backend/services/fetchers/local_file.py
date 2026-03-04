"""Local file fetcher for various document formats."""

import logging
from pathlib import Path
from typing import Any

from langchain_core.documents import Document

from config.settings import settings
from .base import BaseFetcher

logger = logging.getLogger(__name__)


class LocalFileFetcher(BaseFetcher):
    """Fetcher for local files (txt, md, docx, pdf)."""

    def __init__(self, config: dict[str, Any]):
        super().__init__(config)
        
        # Support environment variable substitution
        base_dir = config.get("base_directory", "")
        if base_dir.startswith("${") and base_dir.endswith("}"):
            env_var = base_dir[2:-1]
            base_dir = getattr(settings, env_var, "")
        
        self.base_directory = Path(base_dir) if base_dir else None
        self.supported_extensions = set(config.get("supported_extensions", [".txt", ".md"]))
        self.recursive = config.get("recursive", True)
        self.encoding = config.get("encoding", "utf-8")

    def validate_config(self) -> bool:
        if not settings.KNOWLEDGE_FILES_ENABLED:
            logger.info("LocalFileFetcher: Disabled in settings")
            return False
        
        if not self.base_directory:
            logger.warning("LocalFileFetcher: No base directory configured")
            return False
        
        if not self.base_directory.exists():
            logger.warning(f"LocalFileFetcher: Directory does not exist: {self.base_directory}")
            return False
        
        if not self.base_directory.is_dir():
            logger.error(f"LocalFileFetcher: Path is not a directory: {self.base_directory}")
            return False
        
        return True

    def fetch(self) -> list[Document]:
        """Fetch documents from local files."""
        if not self.validate_config():
            return []

        docs: list[Document] = []
        file_count = 0

        # Use rglob for recursive, glob for non-recursive
        pattern = "**/*" if self.recursive else "*"
        
        for file_path in sorted(self.base_directory.glob(pattern)):
            if not file_path.is_file():
                continue
            
            if file_path.suffix.lower() not in self.supported_extensions:
                continue

            file_count += 1
            try:
                doc = self._load_file(file_path)
                if doc:
                    docs.append(doc)
                    logger.info(f"Loaded local file: {file_path.name}")
            except Exception as e:
                logger.error(f"Failed to load file {file_path}: {e}")

        logger.info(
            f"Local file scan complete: {file_count} files found, "
            f"{len(docs)} documents loaded from {self.base_directory}"
        )
        return docs

    def _load_file(self, file_path: Path) -> Document | None:
        """Load a single file based on its extension."""
        ext = file_path.suffix.lower()
        
        base_metadata = {
            "source": str(file_path),
            "type": "local_file",
            "fetcher": "LocalFileFetcher",
            "format": ext.lstrip("."),
            "filename": file_path.name,
        }

        if ext in (".txt", ".md"):
            return self._load_text_file(file_path, base_metadata)
        elif ext == ".docx":
            return self._load_docx_file(file_path, base_metadata)
        elif ext == ".pdf":
            return self._load_pdf_file(file_path, base_metadata)
        
        return None

    def _load_text_file(self, file_path: Path, metadata: dict) -> Document:
        """Load plain text or markdown file."""
        text = file_path.read_text(encoding=self.encoding, errors="replace")
        return Document(page_content=text, metadata=metadata)

    def _load_docx_file(self, file_path: Path, metadata: dict) -> Document | None:
        """Load Microsoft Word document."""
        try:
            import docx
            
            doc = docx.Document(str(file_path))
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            
            if text.strip():
                return Document(page_content=text, metadata=metadata)
        except ImportError:
            logger.error(
                "python-docx not installed, skipping .docx files. "
                "Run: uv add python-docx"
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {e}")
        
        return None

    def _load_pdf_file(self, file_path: Path, metadata: dict) -> Document | None:
        """Load PDF document."""
        try:
            import fitz  # pymupdf
            
            pdf_doc = fitz.open(str(file_path))
            pages_text = []
            
            for page in pdf_doc:
                page_text = page.get_text()
                if page_text.strip():
                    pages_text.append(page_text)
            
            pdf_doc.close()

            if pages_text:
                text = "\n\n".join(pages_text)
                metadata["page_count"] = len(pages_text)
                return Document(page_content=text, metadata=metadata)
        except ImportError:
            logger.error(
                "pymupdf not installed, skipping .pdf files. "
                "Run: uv add pymupdf"
            )
        except Exception as e:
            logger.error(f"Failed to parse PDF file {file_path}: {e}")
        
        return None
